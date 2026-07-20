"""
文件解析器

职责：把不同格式的文件转成纯文本或结构化元素
支持：PDF、Markdown、TXT、DOCX
"""

from pathlib import Path
from dataclasses import dataclass, field
from typing import Optional
import logging

import chardet

logger = logging.getLogger(__name__)


@dataclass
class DocElement:
    """文档元素，保留结构信息"""
    type: str          # "title" / "paragraph" / "table_header" / "table_row"
    level: int         # 标题层级（1=H1, 2=H2, ...），非标题为 0
    content: str       # 内容（已剥离特殊标记）
    metadata: dict     # 元数据（table_id, warning, info, source_file 等）


def parse_file(file_path: str) -> str:
    """根据文件扩展名选择解析器，返回纯文本"""
    ext = Path(file_path).suffix.lower()

    if ext == ".pdf":
        return parse_pdf(file_path)
    elif ext in (".md", ".markdown"):
        return parse_markdown(file_path)
    elif ext == ".txt":
        return parse_txt(file_path)
    elif ext == ".docx":
        return parse_docx(file_path)
    else:
        raise ValueError(f"不支持的文件格式：{ext}")


def parse_pdf(file_path: str) -> str:
    """解析 PDF，提取文本内容"""
    import fitz  # PyMuPDF
    doc = fitz.open(file_path)
    text_parts = [page.get_text() for page in doc]
    doc.close()
    return "\n".join(text_parts)


def parse_markdown(file_path: str) -> str:
    """解析 Markdown"""
    encoding = get_encoding(file_path)
    return Path(file_path).read_text(encoding=encoding)


def parse_txt(file_path: str) -> str:
    """解析纯文本"""
    encoding = get_encoding(file_path)
    return Path(file_path).read_text(encoding=encoding)


def parse_docx(file_path: str) -> str:
    """解析 Word 文档，返回纯文本（兼容旧接口）"""
    elements = parse_docx_structured(file_path)
    return "\n".join(elem.content for elem in elements)


def get_encoding(file_path: str) -> str:
    """获取文件编码"""
    with open(file_path, 'rb') as f:
        raw_data = f.read()
    detected = chardet.detect(raw_data)
    encoding = detected.get('encoding', 'utf-8') or 'utf-8'
    return encoding


# ========== DOCX 结构化解析 ==========

# 中文标题样式名映射
TITLE_STYLE_MAP = {
    "标题 1": 1, "标题 2": 2, "标题 3": 3,
    "标题1": 1, "标题2": 2, "标题3": 3,
    "Heading 1": 1, "Heading 2": 2, "Heading 3": 3,
    "heading 1": 1, "heading 2": 2, "heading 3": 3,
}

# 话术对关键词
# DIALOGUE_KEYWORDS = ["客户表述", "应对策略", "客户问题", "标准回答", "话术", "回复"]

# 噪声关键词（模板占位符）
NOISE_KEYWORDS = ["序号", "客户昵称", "备注", "日期", "签名", "填写", "待补充", "模板"]


def _extract_special_markers(text: str) -> tuple[str, dict]:
    """
    提取特殊标记，返回 (干净文本, 元数据)

    [!] → warning: True（合规提醒/红线）
    [i] → info: True（话术要点）
    """
    metadata = {}
    clean_text = text.strip()

    if "[!]" in clean_text:
        metadata["warning"] = True
        clean_text = clean_text.replace("[!]", "").strip()
    if "[i]" in clean_text:
        metadata["info"] = True
        clean_text = clean_text.replace("[i]", "").strip()

    return clean_text, metadata


def _get_title_level(style_name: str) -> Optional[int]:
    """
    检测标题样式，返回层级（1/2/3），非标题返回 None

    支持中英文样式名：
    - 英文：Heading 1, Heading 2, Heading 3
    - 中文：标题 1, 标题2, 标题 1
    """
    if not style_name:
        return None

    # 精确匹配
    if style_name in TITLE_STYLE_MAP:
        return TITLE_STYLE_MAP[style_name]

    # 模糊匹配：包含 "Heading" 或 "标题"
    style_lower = style_name.lower()
    if "heading" in style_lower:
        for num in [1, 2, 3]:
            if str(num) in style_lower:
                return num
    if "标题" in style_name:
        for num in [1, 2, 3]:
            if str(num) in style_name:
                return num

    return None


def _is_noise(text: str) -> bool:
    """判断是否为噪声内容"""
    stripped = text.strip()
    if not stripped or len(stripped) < 2:
        return True
    if stripped in NOISE_KEYWORDS:
        return True
    return False


def _get_cell_text(cell) -> str:
    """
    获取单元格文本，处理合并单元格

    合并单元格（gridSpan/vMerge）中，被合并的单元格内容为空
    """
    try:
        text = cell.text.strip()
        # 处理合并单元格：如果单元格有 vMerge 属性且不是第一个，返回空
        tc = cell._tc
        tcPr = tc.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tcPr')
        if tcPr is not None:
            vMerge = tcPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}vMerge')
            if vMerge is not None:
                val = vMerge.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
                # vMerge 没有 val 或 val="continue" 表示是合并的后续单元格
                if val is None or val == "continue":
                    return ""
        return text
    except Exception:
        return cell.text.strip()


def _extract_table_headers(table) -> list[str]:
    """提取表格表头"""
    if not table.rows:
        return []
    return [_get_cell_text(cell) for cell in table.rows[0].cells]


# def _is_dialogue_pair_table(headers: list[str]) -> bool:
#     """检测是否为话术对表格"""
#     return any(
#         any(kw in h for kw in DIALOGUE_KEYWORDS)
#         for h in headers
#     )


def parse_docx_structured(file_path: str) -> list[DocElement]:
    """
    结构化解析 DOCX，返回元素列表

    改进点：
    1. 按文档实际顺序遍历（段落+表格交替）
    2. 支持中英文标题样式
    3. 处理合并单元格
    4. 表格行挂载表头信息
    5. 话术对正确配对
    6. 特殊标记剥离并提取为元数据
    7. 噪声过滤
    8. 添加 source_file 字段
    """
    try:
        from docx import Document
    except ImportError:
        logger.error("python-docx 未安装，无法解析 DOCX")
        return []

    try:
        doc = Document(file_path)
    except Exception as e:
        logger.error(f"DOCX 文件打开失败: {e}")
        return []

    elements = []
    table_counter = 0
    source_file = Path(file_path).name

    # 按文档实际顺序遍历（关键：段落和表格交替出现）
    for child in doc.element.body:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag

        # ===== 处理段落 =====
        if tag == "p":
            try:
                # 找到对应的段落对象
                para = None
                for p in doc.paragraphs:
                    if p._element is child:
                        para = p
                        break

                if para is None:
                    continue

                text = para.text
                if not text or _is_noise(text):
                    continue

                # 提取特殊标记
                clean_text, metadata = _extract_special_markers(text)
                if not clean_text:
                    continue

                metadata["source_file"] = source_file

                # 检测标题
                style_name = para.style.name if para.style else ""
                title_level = _get_title_level(style_name)

                if title_level is not None:
                    elements.append(DocElement(
                        type="title",
                        level=title_level,
                        content=clean_text,
                        metadata=metadata,
                    ))
                else:
                    elements.append(DocElement(
                        type="paragraph",
                        level=0,
                        content=clean_text,
                        metadata=metadata,
                    ))
            except Exception as e:
                logger.warning(f"段落解析异常: {e}")
                continue

        # ===== 处理表格 =====
        elif tag == "tbl":
            try:
                # 找到对应的表格对象
                table = None
                for t in doc.tables:
                    if t._element is child:
                        table = t
                        break

                if table is None:
                    continue

                headers = _extract_table_headers(table)

                # 构建完整的表格内容
                table_rows = []
                for i, row in enumerate(table.rows):
                    if i == 0:
                        continue  # 跳过表头行

                    cells = [_get_cell_text(cell) for cell in row.cells]

                    # 跳过空行
                    if all(not c for c in cells):
                        continue

                    # 剥离特殊标记
                    clean_cells = []
                    for cell_text in cells:
                        clean, _ = _extract_special_markers(cell_text)
                        clean_cells.append(clean)

                    # 键值对格式
                    if headers:
                        kv = ", ".join(f"{h}: {c}" for h, c in zip(headers, clean_cells) if c)
                    else:
                        kv = " | ".join(clean_cells)
                    table_rows.append(kv)

                # 整表作为一个 DocElement
                if table_rows:
                    header_line = ", ".join(headers) if headers else ""
                    content = f"表格[{header_line}]\n" + "\n".join(
                        f"{i+1}. {r}" for i, r in enumerate(table_rows)
                    )
                    elements.append(DocElement(
                        type="table",
                        level=0,
                        content=content,
                        metadata={
                            "table_id": table_counter,
                            "headers": headers,
                            "source_file": source_file,
                        },
                    ))


                table_counter += 1

            except Exception as e:
                logger.warning(f"表格解析异常: {e}")
                continue

    return elements
